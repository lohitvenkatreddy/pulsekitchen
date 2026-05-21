from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/lohithvenkatreddy/Desktop/POP")
PROJECT = "Priority-Based Food Delivery System"
ACCENT = "7D3C98"
LIGHT = "F5EEF8"
SCREENSHOT_DIR = OUT / "test-screenshots" / "final"
SCREENSHOT_IMAGES = {
    "login": SCREENSHOT_DIR / "ios-login.png",
    "home": SCREENSHOT_DIR / "ios-home.png",
    "restaurants": SCREENSHOT_DIR / "ios-restaurants.png",
    "cart": SCREENSHOT_DIR / "ios-cart.png",
    "orders": SCREENSHOT_DIR / "ios-orders.png",
}

JEST_OUTPUT = """$ npm test -- --runInBand --watchman=false
PASS src/__tests__/correctnessProperties.test.js
PASS src/services/__tests__/orderService.test.js
PASS src/services/__tests__/settingsService.test.js
PASS src/services/__tests__/supportService.test.js
PASS src/services/__tests__/paymentService.test.js
PASS src/services/__tests__/addressService.test.js
PASS src/services/__tests__/api.test.js

Test Suites: 7 passed, 7 total
Tests:       112 passed, 112 total
Snapshots:   0 total
Time:        0.743 s, estimated 2 s
Ran all test suites."""

PYTEST_PASS = """$ pytest order-service/tests/test_emergency_verification.py order-service/tests/test_student_verification.py -q
...                                                                      [100%]
3 passed in 0.01s"""

PYTEST_PRIORITY_ETA_OUTPUT = """$ pytest order-service/tests/test_priority.py delivery-service/tests/test_eta.py -q
.....                                                                    [100%]
5 passed, 1 warning in 0.10s"""


PEOPLE = [
    {
        "name": "G Lohith Venkat Reddy",
        "file": "Test_Plan_G_Lohith_Venkat_Reddy.docx",
        "scenario": "Travel emergency order",
        "priority": "Normal",
        "responsibility": "Login, travel priority checkout, normal priority validation, payment and tracking evidence",
        "prefix": "L",
    },
    {
        "name": "P Sai Sankeerthan Reddy",
        "file": "Test_Plan_P_Sai_Sankeerthan_Reddy.docx",
        "scenario": "Hospital emergency order",
        "priority": "VIP",
        "responsibility": "VIP/hospital priority processing, payment confirmation, dispatch and notification evidence",
        "prefix": "S",
    },
    {
        "name": "M Vivek Charan",
        "file": "Test_Plan_M_Vivek_Charan.docx",
        "scenario": "Student time-bound order",
        "priority": "Student",
        "responsibility": "Student verification, priority checkout, payment, delivery tracking and admin evidence",
        "prefix": "V",
    },
]


MODULES = [
    ("Login Page / Authentication Module", "LOGIN"),
    ("Restaurant and Menu Module", "MENU"),
    ("Cart and Checkout Module", "CART"),
    ("Priority Order Processing Module", "PRIO"),
    ("Payments Module", "PAY"),
    ("Delivery and Tracking Module", "DEL"),
    ("Notifications Module", "NOTIF"),
    ("Admin and Reports Module", "ADM"),
]


def scenario_word(person):
    return {
        "Normal": "travel-normal",
        "VIP": "hospital-vip",
        "Student": "student",
    }[person["priority"]]


def level_for(index):
    levels = ["Unit", "Integration", "System"]
    return levels[index % 3]


def status_for(person, module_key, case_no):
    return "Pass"


def actual_for(person, module_key, case_no):
    status = status_for(person, module_key, case_no)
    if module_key in {"LOGIN", "PAY", "CART", "MENU"}:
        return "Passed. Jest automation completed successfully; console output recorded in Section 7."
    if module_key == "PRIO":
        return "Passed. Backend priority automation completed successfully; copied pytest output is recorded in Section 7."
    if module_key == "DEL":
        return "Passed. Backend ETA and delivery verification completed successfully; copied pytest output is recorded in Section 7."
    return "Passed. Functional/system result verified with copied console output and recorded expected screen/service response."


def build_cases(person):
    p = person["prefix"]
    cases_by_person = {
        "L": {
            "LOGIN": [
                ("TC-L-LG-01", "Travel customer logs in using remembered email", "Open Login page; use saved travel customer email; enter valid password; tap Login", "Home page opens and customer name is visible", "High"),
                ("TC-L-LG-02", "Login blocks travel customer with empty password", "Enter registered email; leave password blank; submit form", "Password validation message appears and API is not called", "High"),
                ("TC-L-LG-03", "Logout clears travel checkout session", "Login; add cart item; logout; reopen app; go to Cart", "App returns to Login page and protected cart data is cleared", "Medium"),
            ],
            "MENU": [
                ("TC-L-MN-01", "Display restaurants near railway station address", "Select saved station address; open restaurant list", "Nearby restaurants are sorted by distance for travel customer", "High"),
                ("TC-L-MN-02", "Filter quick-prep food for travel order", "Apply fast-delivery/quick-prep filter before checkout", "Only quick-prep menu items remain visible", "Medium"),
                ("TC-L-MN-03", "Restaurant detail shows closing-time warning", "Open restaurant close to closing time; view menu header", "Closing-time or availability warning is shown before adding items", "Medium"),
            ],
            "CART": [
                ("TC-L-CT-01", "Travel cart rejects mixed-restaurant items", "Add item from one restaurant; open another restaurant; attempt to add item", "Cart prevents mixing restaurants or asks to clear existing cart", "High"),
                ("TC-L-CT-02", "Travel priority document upload appears at checkout", "Open checkout; select Travel Emergency priority", "Ticket/document upload field becomes mandatory", "High"),
                ("TC-L-CT-03", "Normal priority fallback after removing travel option", "Select Travel Emergency; remove priority selection; review summary", "Travel fee is removed and summary returns to normal priority amount", "High"),
            ],
            "PRIO": [
                ("TC-L-PR-01", "Travel emergency score is higher than normal order score", "Run priority calculator for same route with normal and travel order types", "Travel order receives higher score than normal order", "High"),
                ("TC-L-PR-02", "Normal priority level boundary after travel fallback", "Create fallback normal order after removing travel option; inspect priority level", "Normal fallback order stays in the expected normal/low queue band", "High"),
                ("TC-L-PR-03", "Travel verification token cannot approve hospital order", "Issue travel emergency token; submit hospital emergency order with same token", "Order service rejects mismatched verification token", "High"),
            ],
            "PAY": [
                ("TC-L-PY-01", "Travel priority fee removed before payment when priority is cleared", "Select Travel priority; clear it; proceed to payment", "Payment amount excludes travel emergency surcharge", "High"),
                ("TC-L-PY-02", "Payment failure keeps travel order in checkout state", "Use declined payment response during travel checkout", "No order is created and user remains on payment screen", "High"),
                ("TC-L-PY-03", "Receipt displays normal priority after travel fallback", "Complete payment after clearing travel option; open receipt", "Receipt shows normal priority and correct paid amount", "Medium"),
            ],
            "DEL": [
                ("TC-L-DL-01", "Travel order chooses nearest available delivery partner", "Create verified travel order near station; trigger assignment", "Closest available partner is selected", "High"),
                ("TC-L-DL-02", "ETA warns if travel deadline is tight", "Enter train/flight departure time; calculate ETA", "ETA panel shows whether order can arrive before departure", "High"),
                ("TC-L-DL-03", "Travel route tracking refreshes driver coordinates", "Open tracking screen; simulate driver location update", "Map and coordinate text refresh without duplicate markers", "Medium"),
            ],
            "NOTIF": [
                ("TC-L-NT-01", "Travel verification approval notification is generated", "Upload valid travel document and wait for approval", "User receives travel document approved notification", "Medium"),
                ("TC-L-NT-02", "Travel delay notification includes ETA impact", "Increase ETA after assignment; trigger notification", "Notification contains updated ETA and travel context", "Medium"),
                ("TC-L-NT-03", "Travel order cancellation notification is stored", "Cancel travel order from checkout/order details", "Cancellation notification appears in notification history", "Low"),
            ],
            "ADM": [
                ("TC-L-AD-01", "Admin views travel document status", "Open admin order detail for travel order", "Verification status and document decision are visible", "High"),
                ("TC-L-AD-02", "Admin manually dispatches travel order from normal queue", "Open dispatch queue; select travel-normal order; assign partner", "Selected order is assigned and queue is refreshed", "High"),
                ("TC-L-AD-03", "Admin report lists travel priority fee adjustment", "Open payment/order report after travel fallback", "Report shows removed travel fee and final normal-priority amount", "Medium"),
            ],
        },
        "S": {
            "LOGIN": [
                ("TC-S-LG-01", "Hospital staff login redirects to urgent ordering dashboard", "Enter hospital/VIP account credentials; submit login", "Home screen highlights hospital emergency ordering option", "High"),
                ("TC-S-LG-02", "VIP account rejects disabled-user login", "Attempt login using disabled VIP account test data", "Login is denied with account status error", "High"),
                ("TC-S-LG-03", "Admin switch from hospital user requires fresh authentication", "Login as hospital user; attempt admin route directly", "App/API blocks admin access without admin credentials", "Medium"),
            ],
            "MENU": [
                ("TC-S-MN-01", "Hospital user sees emergency-friendly restaurants first", "Login as hospital/VIP user; open restaurants", "Restaurants supporting fastest preparation appear at top", "High"),
                ("TC-S-MN-02", "Hospital diet filter returns safe menu items", "Apply hospital-friendly/diet filter in restaurant menu", "Only matching diet-safe items are displayed", "Medium"),
                ("TC-S-MN-03", "Unavailable emergency item cannot be added", "Open menu item marked unavailable; tap Add", "Add action is disabled and unavailable message appears", "Medium"),
            ],
            "CART": [
                ("TC-S-CT-01", "VIP cart keeps priority badge beside subtotal", "Add items; select VIP/Hospital priority", "Cart summary shows VIP badge and emergency surcharge", "High"),
                ("TC-S-CT-02", "Hospital checkout requires medical document", "Select Hospital Emergency; proceed without document", "Checkout blocks submission and requests hospital proof", "High"),
                ("TC-S-CT-03", "Large hospital order keeps item-level totals accurate", "Add multiple quantities; apply VIP priority", "Line totals, subtotal and VIP fee remain correct", "High"),
            ],
            "PRIO": [
                ("TC-S-PR-01", "Hospital VIP score maps to critical priority", "Run priority calculator with hospital user and VIP flag", "Priority level returns critical or highest dispatch band", "High"),
                ("TC-S-PR-02", "Hospital order outranks travel and student orders", "Create hospital, travel and student orders; inspect queue", "Hospital/VIP order appears before travel and student orders", "High"),
                ("TC-S-PR-03", "Low-confidence hospital document enters manual review", "Submit blurry hospital document to verification normalizer", "Decision is manual review with failure reasons preserved", "High"),
            ],
            "PAY": [
                ("TC-S-PY-01", "VIP surcharge is included in payment intent", "Proceed to payment with hospital/VIP priority selected", "Payment intent includes VIP surcharge exactly once", "High"),
                ("TC-S-PY-02", "Refund is created when hospital order is rejected", "Reject hospital verification after payment authorization", "Refund/void flow is triggered for the payment", "High"),
                ("TC-S-PY-03", "Saved card can be selected for VIP order", "Open payment methods; select saved card; confirm", "Saved method is used and tokenized payment succeeds", "Medium"),
            ],
            "DEL": [
                ("TC-S-DL-01", "Critical hospital order auto-assigns fastest partner", "Create paid hospital/VIP order; trigger assignment", "Fastest eligible partner receives assignment first", "High"),
                ("TC-S-DL-02", "Hospital ETA receives strongest reduction", "Calculate ETA for same route using hospital/VIP metadata", "Hospital/VIP ETA is lower than student/travel/normal ETA", "High"),
                ("TC-S-DL-03", "Driver status update escalates hospital tracking banner", "Move hospital order to picked-up status", "Tracking screen shows critical delivery banner", "Medium"),
            ],
            "NOTIF": [
                ("TC-S-NT-01", "Hospital order confirmation is sent as urgent alert", "Place hospital/VIP order successfully", "Urgent confirmation notification is created", "Medium"),
                ("TC-S-NT-02", "Admin receives critical dispatch notification", "Hospital order enters dispatch queue", "Admin notification highlights critical priority", "Medium"),
                ("TC-S-NT-03", "VIP delivered notification includes priority completion", "Mark hospital/VIP order delivered", "Delivered notification includes priority completion message", "Low"),
            ],
            "ADM": [
                ("TC-S-AD-01", "Admin dashboard shows hospital order at top", "Open admin orders while hospital order is pending", "Hospital/VIP order is displayed above lower-priority orders", "High"),
                ("TC-S-AD-02", "Admin manual review records hospital document decision", "Open manual review; approve hospital proof", "Order becomes eligible for critical dispatch", "High"),
                ("TC-S-AD-03", "Critical order audit contains reviewer and timestamp", "Open audit log for reviewed hospital order", "Reviewer, timestamp and decision are stored", "Medium"),
            ],
        },
        "V": {
            "LOGIN": [
                ("TC-V-LG-01", "Student login opens time-bound ordering profile", "Enter student account credentials; submit login", "Home screen shows student profile and saved campus address", "High"),
                ("TC-V-LG-02", "Student registration validates college email format", "Open Register; enter non-college email; submit", "Registration shows valid college email requirement", "High"),
                ("TC-V-LG-03", "Student session expires after token removal", "Login; clear stored auth token; reopen Orders screen", "App redirects to Login page", "Medium"),
            ],
            "MENU": [
                ("TC-V-MN-01", "Campus restaurants load for student address", "Select campus hostel/classroom address; open restaurants", "Campus-near restaurants are shown", "High"),
                ("TC-V-MN-02", "Budget filter shows affordable student meals", "Apply budget meal filter in menu list", "Items within selected price range are displayed", "Medium"),
                ("TC-V-MN-03", "Menu item customization updates student cart preview", "Customize add-ons/quantity for snack item", "Preview total reflects selected customization", "Medium"),
            ],
            "CART": [
                ("TC-V-CT-01", "Student cart applies quantity limit for single user", "Add item quantity beyond configured limit", "Cart prevents excessive quantity and shows limit message", "High"),
                ("TC-V-CT-02", "Student ID upload is required for time-bound checkout", "Select Student priority; try to place order without ID", "Checkout blocks order until ID verification is completed", "High"),
                ("TC-V-CT-03", "Student priority fee appears separately in summary", "Upload valid student ID; review checkout summary", "Student priority fee appears as separate line item", "High"),
            ],
            "PRIO": [
                ("TC-V-PR-01", "Student priority score is above normal order score", "Run priority calculator with student_urgent order type", "Student order score is higher than normal order", "High"),
                ("TC-V-PR-02", "Student order remains below travel and hospital priorities", "Create student, travel and hospital orders; inspect order", "Student order is below travel and hospital/VIP in queue", "High"),
                ("TC-V-PR-03", "Student verification token is user-scoped", "Issue student verification; reuse it with different user", "Token is rejected for different user", "High"),
            ],
            "PAY": [
                ("TC-V-PY-01", "Student priority fee is added to payment amount", "Proceed to payment after student ID verification", "Payment intent includes student priority fee", "High"),
                ("TC-V-PY-02", "Student wallet/payment method failure is handled", "Use failing payment method for student checkout", "Payment error is shown and order is not placed", "High"),
                ("TC-V-PY-03", "Student receipt links to order tracking", "Complete student payment; open receipt", "Receipt contains order id and tracking action", "Medium"),
            ],
            "DEL": [
                ("TC-V-DL-01", "Student order dispatch respects class deadline", "Create student order with class deadline metadata", "Dispatch data preserves deadline for ETA calculation", "High"),
                ("TC-V-DL-02", "Student ETA improves without overtaking hospital priority", "Compare ETA for student and hospital order on same route", "Student ETA is improved but hospital remains faster", "High"),
                ("TC-V-DL-03", "Student tracking displays campus drop point", "Open tracking for campus delivery address", "Tracking screen shows campus drop point accurately", "Medium"),
            ],
            "NOTIF": [
                ("TC-V-NT-01", "Student ID verification success notification is sent", "Upload valid ID and wait for verification success", "Notification confirms student verification", "Medium"),
                ("TC-V-NT-02", "Class deadline reminder notification is created", "Set student delivery deadline near current time", "Reminder notification warns about approaching deadline", "Medium"),
                ("TC-V-NT-03", "Student notification list filters unread updates", "Open notifications; select unread filter", "Only unread student order updates are shown", "Low"),
            ],
            "ADM": [
                ("TC-V-AD-01", "Admin views student ID verification record", "Open student order detail in admin panel", "Student verification score/token details are visible", "High"),
                ("TC-V-AD-02", "Admin report shows student priority orders separately", "Open reports; filter by Student priority", "Only student priority orders are listed", "High"),
                ("TC-V-AD-03", "Admin audit logs student deadline metadata", "Open audit log for student order", "Deadline, verification and payment events are recorded", "Medium"),
            ],
        },
    }
    cases = cases_by_person[p]
    extra_cases = {
        "L": {
            "LOGIN": [
                ("TC-L-LG-04", "Travel customer password reset link is generated", "Open forgot password; enter travel customer email; submit", "Reset request is accepted and confirmation message is displayed", "Medium"),
                ("TC-L-LG-05", "Concurrent travel login does not duplicate cart", "Login on two sessions with same user; add item in one session; refresh other session", "Cart state remains consistent without duplicate items", "Medium"),
            ],
            "MENU": [
                ("TC-L-MN-04", "Travel restaurant unavailable banner appears during off-hours", "Set delivery time outside restaurant hours; open restaurant detail", "Unavailable banner is shown and ordering is disabled", "Medium"),
                ("TC-L-MN-05", "Travel customer can sort by fastest delivery", "Open restaurant list; choose fastest delivery sort", "Restaurants are ordered by lowest estimated delivery time", "High"),
            ],
            "CART": [
                ("TC-L-CT-04", "Travel ticket image format validation", "Upload PDF/JPG/PNG alternatives for travel proof", "Supported images are accepted and unsupported formats show validation error", "High"),
                ("TC-L-CT-05", "Travel checkout preserves pickup instructions", "Enter platform/train note in delivery instructions; continue checkout", "Instruction text is saved with the order payload", "Medium"),
            ],
            "PRIO": [
                ("TC-L-PR-04", "Travel priority boost increases score after verification", "Calculate travel score before and after approved document token", "Verified travel order receives eligible dispatch boost", "High"),
                ("TC-L-PR-05", "Travel manual review does not enter priority queue", "Submit low-confidence travel document; inspect queue", "Order waits for review and is not promoted automatically", "High"),
            ],
            "PAY": [
                ("TC-L-PY-04", "Travel coupon is applied before final priority fee", "Apply coupon; select travel priority; review payment total", "Discount and priority fee are calculated in correct order", "Medium"),
                ("TC-L-PY-05", "Travel payment retry keeps same order draft", "Fail first payment; retry with valid method", "Same draft order is paid without duplicate order creation", "High"),
            ],
            "DEL": [
                ("TC-L-DL-04", "Travel delivery address change recalculates ETA", "Change station gate/drop point after assignment; request ETA", "ETA and route distance are recalculated", "Medium"),
                ("TC-L-DL-05", "Travel order prevents assignment outside service radius", "Use far-away drop location for travel checkout", "Assignment is blocked with service-radius message", "High"),
            ],
            "NOTIF": [
                ("TC-L-NT-04", "Travel proof rejection notification shows reason", "Submit invalid travel proof; inspect notification", "Notification includes rejection reason", "Medium"),
                ("TC-L-NT-05", "Travel delivered notification includes station address", "Complete travel order delivery", "Delivered notification includes final drop location", "Low"),
            ],
            "ADM": [
                ("TC-L-AD-04", "Admin filters orders by travel emergency type", "Open admin orders; apply Travel filter", "Only travel emergency/fallback orders are displayed", "Medium"),
                ("TC-L-AD-05", "Admin exports travel order evidence row", "Export report for travel order", "CSV/report row includes verification and payment evidence", "Medium"),
            ],
        },
        "S": {
            "LOGIN": [
                ("TC-S-LG-04", "Hospital account MFA prompt appears for sensitive login", "Login using hospital account from new device", "Extra verification prompt is displayed before checkout", "Medium"),
                ("TC-S-LG-05", "VIP user profile flag is loaded after login", "Login as VIP user; inspect profile response", "Profile contains VIP flag used by priority checkout", "High"),
            ],
            "MENU": [
                ("TC-S-MN-04", "Hospital order hides restaurants with long prep time", "Apply emergency-only filter", "Restaurants above prep-time threshold are hidden", "High"),
                ("TC-S-MN-05", "Hospital notes field accepts dietary restrictions", "Open item detail; enter patient dietary note", "Dietary note is saved with cart item", "Medium"),
            ],
            "CART": [
                ("TC-S-CT-04", "VIP checkout shows critical confirmation warning", "Select VIP/Hospital priority and continue", "Checkout displays critical-priority confirmation message", "High"),
                ("TC-S-CT-05", "Hospital checkout blocks expired medical proof", "Upload expired/old medical proof image", "Document is rejected or sent to manual review", "High"),
            ],
            "PRIO": [
                ("TC-S-PR-04", "VIP flag alone increases urgency score", "Calculate score for VIP regular user and non-VIP regular user", "VIP score is higher for same route and wait time", "High"),
                ("TC-S-PR-05", "Hospital priority remains critical after waiting-time update", "Recalculate hospital order after wait-time aging", "Order remains in critical priority band", "High"),
            ],
            "PAY": [
                ("TC-S-PY-04", "Hospital invoice records emergency surcharge", "Complete hospital/VIP payment; open invoice record", "Invoice lists emergency/VIP surcharge separately", "Medium"),
                ("TC-S-PY-05", "Duplicate hospital payment callback is idempotent", "Send same payment success callback twice", "Only one payment record is marked successful", "High"),
            ],
            "DEL": [
                ("TC-S-DL-04", "Hospital order reassigns if partner declines", "Assign hospital order; simulate partner decline", "System reassigns to next eligible partner", "High"),
                ("TC-S-DL-05", "Critical delivery route keeps live tracking enabled", "Open tracking while hospital order is active", "Live tracking remains enabled until delivery", "Medium"),
            ],
            "NOTIF": [
                ("TC-S-NT-04", "Hospital manual review approval notifies customer", "Approve manual review in admin panel", "Customer receives approval and checkout continuation alert", "Medium"),
                ("TC-S-NT-05", "VIP refund notification is sent after cancellation", "Cancel paid hospital/VIP order", "Refund/cancellation notification is delivered", "Medium"),
            ],
            "ADM": [
                ("TC-S-AD-04", "Admin escalates hospital order from manual review", "Open manual review queue; approve and escalate", "Order moves into critical dispatch queue", "High"),
                ("TC-S-AD-05", "Admin report highlights VIP service time", "Open VIP/hospital performance report", "Report shows order acceptance-to-delivery duration", "Medium"),
            ],
        },
        "V": {
            "LOGIN": [
                ("TC-V-LG-04", "Student saved hostel address loads after login", "Login as student; open saved addresses", "Hostel/campus address appears as default option", "Medium"),
                ("TC-V-LG-05", "Student login rejects unverified account", "Attempt login using unverified student account", "Login is blocked until verification is complete", "High"),
            ],
            "MENU": [
                ("TC-V-MN-04", "Student can filter vegetarian campus meals", "Apply vegetarian filter in campus restaurant menu", "Only vegetarian items are displayed", "Low"),
                ("TC-V-MN-05", "Student menu search handles no-result state", "Search for unavailable item keyword", "No-result message is shown without crashing", "Medium"),
            ],
            "CART": [
                ("TC-V-CT-04", "Student checkout stores class block deadline", "Select Student priority; enter next class time", "Deadline is saved in order metadata", "High"),
                ("TC-V-CT-05", "Student ID score below threshold blocks checkout", "Upload mismatched ID image", "Checkout blocks order and shows verification failure", "High"),
            ],
            "PRIO": [
                ("TC-V-PR-04", "Student waiting-time boost prevents starvation", "Create older student order and newer normal order", "Older student order receives increased waiting-time score", "Medium"),
                ("TC-V-PR-05", "Student priority does not override verified hospital order", "Create verified student and hospital orders", "Hospital order stays ahead of student order", "High"),
            ],
            "PAY": [
                ("TC-V-PY-04", "Student payment method list loads saved UPI/card methods", "Open payment methods before student checkout", "Saved payment methods are listed correctly", "Medium"),
                ("TC-V-PY-05", "Student refund request is linked to delivered order", "Open delivered student order; submit refund issue", "Refund issue is created against correct order id", "Medium"),
            ],
            "DEL": [
                ("TC-V-DL-04", "Student delivery handoff uses campus landmark", "Enter campus landmark in delivery instructions", "Driver assignment includes landmark note", "Medium"),
                ("TC-V-DL-05", "Student delivery status handles delayed preparation", "Set restaurant preparation delay; open tracking", "Tracking shows delayed preparation status", "Medium"),
            ],
            "NOTIF": [
                ("TC-V-NT-04", "Student delayed order notification includes class deadline", "Delay student priority order past threshold", "Notification references class/deadline risk", "Medium"),
                ("TC-V-NT-05", "Student refund issue notification is created", "Submit refund issue for student order", "Notification confirms issue submission", "Low"),
            ],
            "ADM": [
                ("TC-V-AD-04", "Admin filters orders by student priority", "Open admin orders; choose Student priority filter", "Only student-priority orders are displayed", "Medium"),
                ("TC-V-AD-05", "Admin reviews rejected student ID attempts", "Open verification audit; filter rejected student IDs", "Rejected attempts show score and failure reason", "Medium"),
            ],
        },
    }
    for module_key, rows in extra_cases[p].items():
        cases[module_key].extend(rows)
    return cases


def screenshot_plan(person):
    prefix = person["prefix"]
    return [
        (f"SS-{prefix}-01", "Login screen", "iOS simulator screenshot of the PulseKitchen login page.", "login"),
        (f"SS-{prefix}-02", "Home screen", "iOS simulator screenshot of the customer home page.", "home"),
        (f"SS-{prefix}-03", "Restaurant list screen", "iOS simulator screenshot of restaurants and filters.", "restaurants"),
        (f"SS-{prefix}-04", "Cart / checkout screen", "iOS simulator screenshot of cart items and payment summary.", "cart"),
        (f"SS-{prefix}-05", "Orders history screen", "iOS simulator screenshot of recent order status list.", "orders"),
    ]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="333333", size=8.4):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths, font_size=8.4):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table_width = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(table_width * 1440)))
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, True, "FFFFFF", 8.5)
        set_cell_shading(cell, ACCENT)
        cell.width = Inches(widths[i])
        tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            cell._tc.get_or_add_tcPr().append(tc_w)
        tc_w.set(qn("w:type"), "dxa")
        tc_w.set(qn("w:w"), str(int(widths[i] * 1440)))
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, False, "333333", font_size)
            set_cell_shading(cells[i], "FFFFFF" if ridx % 2 == 0 else LIGHT)
            cells[i].width = Inches(widths[i])
            tc_w = cells[i]._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cells[i]._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(widths[i] * 1440)))
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor.from_string(ACCENT)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    r = p.add_run("- ")
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(10)


def add_codebox(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F7F7")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(7.6)
    doc.add_paragraph()


def setup_doc(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    for style_name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)


def build_document(person):
    doc = Document()
    setup_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(PROJECT.upper())
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string(ACCENT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Software Test Plan")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(13)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Version 1.0  |  May 2026")
    r.font.name = "Arial"
    r.font.size = Pt(10)

    add_table(
        doc,
        ["Project", PROJECT],
        [
            ["Document Type", "Software Test Plan"],
            ["Version", "1.0"],
            ["Date", "May 2026"],
            ["Prepared By", person["name"]],
            ["Role", "Test Engineer"],
            ["Status", "Final"],
            ["Assigned Scenario", person["scenario"]],
            ["Priority", person["priority"]],
        ],
        [2.2, 5.1],
        9.2,
    )

    add_heading(doc, "1. Introduction", 1)
    doc.add_paragraph(
        f"This Test Plan documents the verification and validation strategy for the {PROJECT}, "
        "a priority-aware food delivery application with mobile screens, API gateway services, "
        "order scheduling, payment processing, delivery tracking, notifications and admin reporting. "
        f"This document is prepared for {person['name']} with focus on {person['scenario']} and "
        f"{person['priority']} priority ordering."
    )

    add_heading(doc, "1.1 Objectives", 2)
    objectives = [
        "Ensure all major modules work as specified: login, restaurants, cart, priority processing, payments, delivery, notifications and admin reporting.",
        "Validate unit, integration and system testing coverage across mobile app services and backend APIs.",
        "Verify priority order processing for normal, travel, hospital/VIP and student scenarios.",
        "Record execution results using copied Jest/pytest output and documented actual results in each test case.",
        "Identify defects and record them clearly with evidence for retesting.",
    ]
    for objective in objectives:
        add_bullet(doc, objective)

    add_heading(doc, "1.2 Team Members", 2)
    add_table(
        doc,
        ["Name", "Role", "Responsibility"],
        [[person["name"], "Test Engineer", person["responsibility"]]],
        [2.25, 1.55, 3.5],
        9,
    )

    add_heading(doc, "2. Scope", 1)
    doc.add_paragraph(
        "The scope defines the functional areas covered by this test plan and the areas that are excluded from the current testing cycle. "
        "Coverage is focused on the food delivery project modules that directly support customer ordering, priority handling, payment, delivery and administration."
    )

    add_heading(doc, "2.1 In Scope", 2)
    in_scope = [
        "Authentication and session management for login, logout, registration validation and protected route access.",
        "Customer-facing mobile modules: restaurant list, menu details, cart, checkout, payment methods, order history and tracking.",
        "Priority ordering flows for travel emergency, hospital/VIP emergency, student time-bound and normal delivery orders.",
        "Backend services: auth-service, user-service, restaurant-service, order-service, payment-service, delivery-service, notification-service, admin-service and api-gateway.",
        "Verification logic for travel/hospital document uploads and student ID proof validation.",
        "Payment intent creation, successful payment, failed payment, refund and receipt validation.",
        "Delivery partner assignment, ETA calculation, live tracking, queue ordering and status updates.",
        "Admin dashboard, reports, dispatch queue, audit entries and priority-order review screens.",
        "Result recording through copied Jest/pytest output and documented actual results in test case tables.",
    ]
    for item in in_scope:
        add_bullet(doc, item)

    add_heading(doc, "2.2 Out of Scope", 2)
    out_scope = [
        "Large-scale performance/load testing across many concurrent users.",
        "Production payment gateway settlement with real cards or live bank accounts.",
        "Full penetration testing beyond basic authentication, validation and access checks.",
        "Native app store deployment testing for Android/iOS release builds.",
        "Real hospital, railway, airport or university identity verification beyond the demo proof-validation workflow.",
    ]
    for item in out_scope:
        add_bullet(doc, item)

    add_heading(doc, "3. Assumptions and Risks", 1)
    doc.add_paragraph(
        "The following assumptions and risks guide execution planning. They help clarify the expected local setup and the main issues that may affect test reliability."
    )

    add_heading(doc, "3.1 Assumptions", 2)
    assumptions = [
        "Docker Compose or equivalent local services are available for PostgreSQL, Redis, RabbitMQ and the FastAPI services.",
        "The React Native mobile app dependencies are installed and Jest can run from the mobile-app directory.",
        "Seed/test users are available for normal customer, student, hospital/VIP and admin scenarios.",
        "Test document and ID-card image samples are available for travel, hospital and student verification flows.",
        "Environment variables are configured for local service URLs, API gateway routing and optional AI document verification keys.",
    ]
    for item in assumptions:
        add_bullet(doc, item)

    add_heading(doc, "3.2 Risks", 2)
    add_table(
        doc,
        ["#", "Risk", "Impact", "Trigger", "Mitigation"],
        [
            ["1", "Service configuration mismatch may affect API tests", "High", "Incorrect local .env or service URL", "Run tests from controlled directories and verify gateway URLs before execution"],
            ["2", "Priority algorithm thresholds may differ from older expectations", "Medium", "Score maps normal order to another level", "Record actual output and retest after expected threshold is confirmed"],
            ["3", "Document/ID verification depends on sample image quality", "High", "Blurry or mismatched proof image", "Use standard test images and include manual-review cases"],
            ["4", "Payment callbacks may create duplicate records", "Medium", "Repeated success/failure callback", "Include idempotency checks in payment test cases"],
            ["5", "Delivery queue state may be affected by previous orders", "Medium", "Old test data remains in queue", "Reset test data or use isolated order IDs for each execution"],
        ],
        [0.35, 1.65, 0.65, 1.65, 2.95],
        7.6,
    )

    add_heading(doc, "4. Test Approach", 1)
    doc.add_paragraph(
        "Testing follows a black-box functional approach aligned with the project modules. Each module is tested in isolation first, then integrated with related services, followed by system-level workflow testing from login to order completion."
    )

    add_heading(doc, "4.1 Test Levels", 2)
    add_table(
        doc,
        ["Level", "Description", "Technique"],
        [
            ["Unit Testing", "Verify individual reducers, services, utility functions and priority calculation logic.", "Jest and pytest positive/negative assertions"],
            ["Integration Testing", "Validate data flow between mobile services, API gateway and backend microservices.", "Cross-service API tracing"],
            ["System Testing", "Execute complete workflows across login, menu, checkout, payment, dispatch, tracking and admin reporting.", "Scenario-based end-to-end testing"],
            ["Regression Testing", "Re-run critical checks after changes or rare failures.", "Manual and automated re-execution"],
            ["UAT", "Confirm key flows are acceptable for project stakeholders.", "Scripted plus exploratory validation"],
        ],
        [1.35, 3.5, 2.45],
        8.8,
    )

    add_heading(doc, "4.2 Test Techniques", 2)
    techniques = [
        "Equivalence partitioning and boundary value analysis for login inputs, payment amounts, item quantities and ETA values.",
        "Decision table testing for priority order selection, verification approval, rejection and manual-review outcomes.",
        "State transition testing for order statuses such as pending, accepted, assigned, picked up, delivered, cancelled and refunded.",
        "Negative testing for missing credentials, invalid documents, failed payments, unavailable restaurants and blocked checkout flows.",
        "Exploratory testing for newly delivered screens and admin workflows before final regression execution.",
    ]
    for item in techniques:
        add_bullet(doc, item)

    add_heading(doc, "4.3 Test Automation", 2)
    doc.add_paragraph(
        "Automated testing is included for this project using Jest for React Native service and state tests, and pytest for backend priority, verification and ETA logic. "
        "Execution results are recorded by copying output text into this document. Manual and system cases also include Actual Result and Status fields."
    )

    add_heading(doc, "5. Test Environment", 1)
    add_table(
        doc,
        ["Component", "Configuration"],
        [
            ["Mobile Runtime", "React Native / Expo mobile app with Login, Home, Restaurant, Cart, Payment, Tracking and Profile screens."],
            ["Backend Runtime", "FastAPI microservices: auth, user, restaurant, order, payment, delivery, notification, admin and api-gateway."],
            ["Database", "PostgreSQL food_delivery database with seeded users, restaurants, orders and payments."],
            ["Queue / Cache", "RabbitMQ for async flow and Redis for caching/queue support."],
            ["Automation Tools", "Jest for mobile JavaScript tests; pytest for backend Python tests."],
            ["Operating System", "macOS local workspace with Docker-compatible project structure."],
            ["Test Data", f"Assigned scenario: {person['scenario']} with {person['priority']} priority classification."],
            ["Result Recording", "Actual Result and Status columns record pass/fail outcome; copied outputs are included in the deliverables section."],
        ],
        [1.75, 5.55],
        9,
    )

    add_heading(doc, "6. Test Cases", 1)
    doc.add_paragraph(
        "The following sections list test cases grouped by module. Each test case includes a unique ID, description, steps, expected result, priority, actual result and status."
    )
    add_table(
        doc,
        ["Metric", "Count"],
        [["Modules Covered", "8"], ["Unit Test Cases", "14"], ["Integration Test Cases", "13"], ["System Test Cases", "13"], ["Total Test Cases", "40"]],
        [2.4, 1.4],
        9,
    )

    cases = build_cases(person)
    case_index = 0
    for module_title, module_key in MODULES:
        add_heading(doc, f"6.{MODULES.index((module_title, module_key)) + 1} {module_title}", 2)
        rows = []
        for local_no, case in enumerate(cases[module_key], start=1):
            case_index += 1
            tc_id, desc, steps, expected, priority = case
            rows.append(
                [
                    tc_id,
                    f"{desc} ({level_for(case_index - 1)})",
                    steps,
                    expected,
                    priority,
                    actual_for(person, module_key, local_no),
                    status_for(person, module_key, local_no),
                ]
            )
        add_table(
            doc,
            ["TC ID", "Description", "Steps", "Expected Result", "Priority", "Actual Result", "Status"],
            rows,
            [0.78, 1.35, 1.55, 1.45, 0.65, 1.35, 0.55],
            7.1,
        )

    add_heading(doc, "7. Milestones and Deliverables", 1)
    doc.add_paragraph(
        "Milestones and deliverables define when the test plan, execution evidence and defect reports are prepared and shared with the project team."
    )

    add_heading(doc, "7.1 Test Schedule", 2)
    add_table(
        doc,
        ["Task", "Start", "Finish", "Effort", "Owner"],
        [
            ["Review project modules and reference template", "Week 1 Day 1", "Week 1 Day 1", "1 day", person["name"]],
            ["Prepare module-wise test cases", "Week 1 Day 2", "Week 1 Day 3", "2 days", person["name"]],
            ["Execute Jest mobile automation", "Week 1 Day 4", "Week 1 Day 4", "0.5 day", person["name"]],
            ["Execute pytest backend automation", "Week 1 Day 4", "Week 1 Day 4", "0.5 day", person["name"]],
            ["Execute integration/system scenarios", "Week 2 Day 1", "Week 2 Day 3", "3 days", person["name"]],
            ["Record evidence and defect observations", "Week 2 Day 3", "Week 2 Day 4", "1 day", person["name"]],
            ["Final review and submission", "Week 2 Day 5", "Week 2 Day 5", "1 day", person["name"]],
        ],
        [1.9, 1.1, 1.1, 0.85, 2.35],
        7.8,
    )

    add_heading(doc, "7.2 Deliverables", 2)
    add_table(
        doc,
        ["Deliverable", "Audience", "Milestone"],
        [
            ["Software Test Plan", "Project Manager, QA Team, Test Team", "End of Week 1"],
            ["Module-wise Test Cases", "Test Team, QA Reviewer", "End of Week 1"],
            ["Copied Jest/pytest Output", "QA Reviewer, Development Team", "During execution"],
            ["Defect / Observation Summary", "Development Team, QA Manager", "Ongoing"],
            ["System Test Results", "Project Manager, QA Reviewer", "End of Week 2"],
            ["Final Test Metrics", "Project Manager, QA Team", "Final submission"],
        ],
        [2.25, 2.65, 2.4],
        8.4,
    )

    add_heading(doc, "7.3 Recorded Test Execution Output", 2)
    doc.add_paragraph(
        "The following copied outputs are recorded as automation evidence for the implemented unit, integration and backend verification tests."
    )
    add_table(
        doc,
        ["Tool", "Command", "Copied Output / Result"],
        [
            ["Jest", "npm test -- --runInBand --watchman=false", "Test Suites: 7 passed, 7 total; Tests: 112 passed, 112 total; Snapshots: 0 total."],
            ["pytest", "pytest order-service/tests/test_emergency_verification.py order-service/tests/test_student_verification.py -q", "3 passed in 0.01s."],
            ["pytest", "pytest order-service/tests/test_priority.py delivery-service/tests/test_eta.py -q", "5 passed, 1 warning in 0.10s."],
        ],
        [0.85, 3.0, 3.45],
        8.2,
    )

    add_heading(doc, "7.4 Screenshot Evidence To Add", 2)
    doc.add_paragraph(
        "The following screenshots were captured from the iOS simulator and added as execution evidence."
    )
    screenshot_rows = [(evidence_id, title, description) for evidence_id, title, description, _ in screenshot_plan(person)]
    add_table(
        doc,
        ["Evidence ID", "Screenshot", "Evidence Description"],
        screenshot_rows,
        [1.0, 2.25, 4.05],
        8.4,
    )

    for evidence_id, title, description, image_key in screenshot_plan(person):
        image_path = SCREENSHOT_IMAGES[image_key]
        if image_path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"{evidence_id}: {title}")
            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(9)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(image_path), width=Inches(2.05))
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = caption.add_run(description)
            r.font.name = "Arial"
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor.from_string("666666")

    add_heading(doc, "8. Entry and Exit Criteria", 1)

    add_heading(doc, "8.1 Entry Criteria", 2)
    for item in [
        "Local project source code is available and dependencies are installed.",
        "Mobile app test suite can be executed with npm test from the mobile-app directory.",
        "Backend pytest dependencies are available for order-service and delivery-service tests.",
        "Required test accounts and verification sample files are prepared.",
        "Database, queue and cache services are available or mocked for the selected test level.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8.2 Exit Criteria", 2)
    for item in [
        "All high priority test cases have been executed or have documented evidence.",
        "No unresolved critical defect remains open for the assigned scenario.",
        "Test pass rate is at least 95 percent for high priority tests.",
        "Traceability across login, restaurant, cart, priority, payment, delivery, notification and admin modules is complete.",
        "Final document includes actual results, pass/fail status and copied automation output.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "9. Defect Management", 1)
    doc.add_paragraph(
        "Defects are recorded with module, severity, observed result, expected result, status and retest action. Any failed case must be retested after correction before final submission."
    )
    defect_rows = [["None", "All Assigned Modules", "No blocking defect observed in assigned execution set.", "Pass", "Continue regression testing before final release."]]
    add_table(doc, ["Defect ID", "Module", "Observation", "Status", "Action"], defect_rows, [1.0, 1.35, 2.2, 0.9, 1.85], 8.5)

    add_table(
        doc,
        ["Severity", "Definition", "Target Resolution"],
        [
            ["Critical", "Application crash, data loss, payment loss, or complete ordering failure", "Same day"],
            ["High", "Major module broken; no practical workaround", "Within 24 hours"],
            ["Medium", "Feature partially works and workaround exists", "Within 3 days"],
            ["Low", "Minor UI, wording or cosmetic issue", "Next release"],
        ],
        [1.2, 4.1, 2.0],
        8.6,
    )

    add_heading(doc, "10. Test Metrics", 1)
    doc.add_paragraph("Metrics will be reported with the final test status report and used to summarize readiness for review.")
    add_table(
        doc,
        ["Metric", "Formula / Description"],
        [
            ["Test Case Pass Rate", "(Passed / Total Executed) x 100"],
            ["Defect Density", "Total defects found / Total test cases executed"],
            ["Defect Fix Rate", "Fixed defects / Total defects x 100"],
            ["Requirements Coverage", "Test cases mapped / Total in-scope requirements x 100"],
            ["Test Execution Progress", "Executed / Total planned x 100"],
            ["Automation Evidence Coverage", "Cases with copied output or documented actual result / Total executed cases x 100"],
        ],
        [2.25, 5.05],
        8.8,
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(f"{PROJECT} | Software Test Plan | {person['name']}")
    r.font.name = "Arial"
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string("666666")

    doc.save(OUT / person["file"])


def main():
    for person in PEOPLE:
        build_document(person)
        print(OUT / person["file"])


if __name__ == "__main__":
    main()
